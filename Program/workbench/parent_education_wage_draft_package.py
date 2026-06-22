from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import RGBColor


SCHEMA_VERSION = "p3.parent_education_wage_draft_package.v1"
TOPIC = "父母受教育水平对子女工资收入的影响"
TOPIC_SLUG = "parent-education-wage"

DEFAULT_P2_READINESS_PATH = Path("Results/json/parent_education_wage_p2_execution_readiness.json")
DEFAULT_PACKAGE_PATH = Path("Results/json/parent_education_wage_p3_draft_package.json")
DEFAULT_MARKDOWN_PATH = Path("Manuscripts/generated/parent_education_wage_paper_draft.md")
DEFAULT_DOCX_PATH = Path("Submissions/parent_education_wage_paper_draft.docx")
DEFAULT_ISSUE_LIST_PATH = Path("Manuscripts/generated/parent_education_wage_issue_list.md")
DEFAULT_AUDIT_REPORT_PATH = Path("Reviews/parent_education_wage_draft_audit_report.md")

FORMAL_VARIABLE_ROLES_PATH = Path("state/product/variable_roles.json")
FORMAL_DESIGN_SPEC_PATH = Path("state/product/design_spec.json")
FORMAL_RUN_PLAN_PATH = Path("state/product/run_plan.json")


def run_parent_education_wage_draft_package(project_root: Path) -> tuple[dict[str, Any], Path]:
    package = build_parent_education_wage_draft_package(project_root)
    write_parent_education_wage_draft_package(project_root, package)
    return package, project_root / DEFAULT_PACKAGE_PATH


def build_parent_education_wage_draft_package(project_root: Path) -> dict[str, Any]:
    readiness = load_json(project_root / DEFAULT_P2_READINESS_PATH)
    blocking_reasons = normalize_list(readiness.get("blocking_reasons"))
    execution_allowed = bool(readiness.get("execution_preflight_allowed"))
    missing_fields = build_missing_field_issues(readiness)
    issue_count = len(blocking_reasons) + len(missing_fields)
    full_draft_ready = execution_allowed and issue_count == 0
    status = "full_draft_package_ready" if full_draft_ready else "blocked_draft_package_ready"
    draft_kind = "full_evidence_bound_draft" if full_draft_ready else "partial_red_flagged_draft"
    topic = str(readiness.get("topic") or TOPIC)
    topic_slug = str(readiness.get("topic_slug") or TOPIC_SLUG)
    return {
        "schema_version": SCHEMA_VERSION,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "topic": topic,
        "topic_slug": topic_slug,
        "status": status,
        "draft_kind": draft_kind,
        "full_draft_ready": full_draft_ready,
        "run_id": readiness.get("run_id"),
        "blocking_reasons": blocking_reasons,
        "issue_count": issue_count,
        "issues": build_issues(readiness, blocking_reasons, missing_fields),
        "source_artifacts": {
            "p2_execution_readiness": {
                "path": DEFAULT_P2_READINESS_PATH.as_posix(),
                "status": readiness.get("status", "missing"),
            }
        },
        "outputs": {
            "json": DEFAULT_PACKAGE_PATH.as_posix(),
            "markdown": DEFAULT_MARKDOWN_PATH.as_posix(),
            "docx": DEFAULT_DOCX_PATH.as_posix(),
            "issue_list": DEFAULT_ISSUE_LIST_PATH.as_posix(),
            "audit_report": DEFAULT_AUDIT_REPORT_PATH.as_posix(),
        },
        "boundary_flags": {
            "modified_formal_variable_roles": False,
            "modified_formal_design_spec": False,
            "modified_formal_run_plan": False,
            "executed_regression": False,
            "created_run_id": False,
            "wrote_full_draft_claims": full_draft_ready,
        },
        "formal_state": {
            "variable_roles": {"path": FORMAL_VARIABLE_ROLES_PATH.as_posix(), "modified": False},
            "design_spec": {"path": FORMAL_DESIGN_SPEC_PATH.as_posix(), "modified": False},
            "run_plan": {"path": FORMAL_RUN_PLAN_PATH.as_posix(), "modified": False},
        },
        "product_control_signal": {
            "phase": "P3",
            "label": "DraftPackage",
            "status": status,
            "primary_artifact": DEFAULT_DOCX_PATH.as_posix(),
            "next_action": "review_partial_draft_and_issue_list" if not full_draft_ready else "review_full_draft_package",
        },
    }


def write_parent_education_wage_draft_package(project_root: Path, package: dict[str, Any]) -> None:
    outputs = package["outputs"]
    markdown_path = project_root / outputs["markdown"]
    docx_path = project_root / outputs["docx"]
    issue_path = project_root / outputs["issue_list"]
    audit_path = project_root / outputs["audit_report"]
    json_path = project_root / outputs["json"]

    markdown_path.parent.mkdir(parents=True, exist_ok=True)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    issue_path.parent.mkdir(parents=True, exist_ok=True)
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.parent.mkdir(parents=True, exist_ok=True)

    markdown_path.write_text(render_markdown_draft(package), encoding="utf-8")
    issue_path.write_text(render_issue_list(package), encoding="utf-8")
    audit_path.write_text(render_audit_report(package), encoding="utf-8")
    write_docx_draft(docx_path, package)
    json_path.write_text(json.dumps(package, ensure_ascii=False, indent=2), encoding="utf-8")


def build_missing_field_issues(readiness: dict[str, Any]) -> list[dict[str, str]]:
    issues = []
    for item in readiness.get("field_supplementation", []):
        if not isinstance(item, dict):
            continue
        status = str(item.get("supplement_status", ""))
        if status not in {"missing", "candidate_found"}:
            continue
        field = str(item.get("dataset_column", ""))
        if field not in {"father_education", "mother_education", "parent_education", "hukou"}:
            continue
        issue_type = "missing_field" if status == "missing" else "candidate_needs_binding"
        issues.append(
            {
                "id": f"{issue_type}:{field}",
                "field": field,
                "status": status,
                "summary": f"{field} 尚未形成正式字段绑定",
                "next_action": "locate_source_field_or_adjust_research_scope"
                if status == "missing"
                else "human_bind_candidate_field",
            }
        )
    return issues


def build_issues(
    readiness: dict[str, Any],
    blocking_reasons: list[str],
    missing_fields: list[dict[str, str]],
) -> list[dict[str, str]]:
    issues = [
        {
            "id": f"blocking_reason:{reason}",
            "type": "blocking_reason",
            "summary": reason,
            "next_action": next_action_for_reason(reason),
        }
        for reason in blocking_reasons
    ]
    issues.extend(missing_fields)
    if not readiness:
        issues.append(
            {
                "id": "missing:p2_execution_readiness",
                "type": "missing_source",
                "summary": "P2 执行准入账本不存在",
                "next_action": "run_p2_execution_readiness_first",
            }
        )
    return issues


def next_action_for_reason(reason: str) -> str:
    if reason == "missing_parent_education_fields":
        return "定位并绑定 father_education / mother_education / parent_education"
    if reason == "human_variable_operationalization_required":
        return "确认父母教育变量合成口径"
    return "人工审阅阻断原因"


def render_markdown_draft(package: dict[str, Any]) -> str:
    return "\n".join(
        [
            f"# {package['topic']}",
            "",
            "## 摘要",
            "",
            "本文拟研究父母受教育水平对子女工资收入的影响。当前版本是半成品论文初稿：研究问题、变量口径和方法路径已经形成草稿，但关键父母教育字段尚未完成真实绑定，因此不能报告回归结果或因果结论。",
            "",
            "【红标】父母教育字段尚未绑定，本文不能进入完整实证结果写作。",
            "",
            "## 研究问题",
            "",
            "本文关注父母受教育水平是否以及如何影响子女工资收入。直觉机制包括家庭教育资源、早期人力资本积累、社会资本和代际机会传递。",
            "",
            "## 数据与变量",
            "",
            "当前可写部分：工资收入、子女教育、年龄、性别、城乡等变量已有候选口径。当前不能写成正式变量表的部分：父亲教育、母亲教育或父母教育综合指标仍缺真实字段绑定。",
            "",
            "【红标】缺少 `father_education`、`mother_education` 或 `parent_education` 的正式字段来源。",
            "",
            "## 识别策略",
            "",
            "在字段补齐前，本文只能保留方法草案。可选路径包括 OLS 基准、父母教育分组比较、稳健性控制，以及在有合格工具变量或外生冲击时进一步讨论 IV/DID。当前版本不执行 IV、DID 或 DML。",
            "",
            "【红标】未执行回归，不能写主结果结论。",
            "",
            "## 预期贡献",
            "",
            "如果字段补齐，本文可以把代际教育传递与劳动收入结果连接起来，形成一个适合本科论文的实证研究框架。",
            "",
            "## 当前问题清单",
            "",
            *[f"- {issue['id']}：{issue['summary']}；下一步：{issue['next_action']}" for issue in package["issues"]],
            "",
        ]
    )


def render_issue_list(package: dict[str, Any]) -> str:
    lines = [
        "# P3 DraftPackage 问题清单",
        "",
        f"- 题目：{package['topic']}",
        f"- 状态：`{package['status']}`",
        f"- draft_kind：`{package['draft_kind']}`",
        f"- full_draft_ready：{str(package['full_draft_ready']).lower()}",
        "",
        "## 阻断项",
    ]
    lines.extend(f"- `{reason}`" for reason in package["blocking_reasons"] or ["none"])
    lines.extend(["", "## 具体问题"])
    lines.extend(f"- `{issue['id']}`：{issue['summary']}；下一步：{issue['next_action']}" for issue in package["issues"])
    lines.append("")
    return "\n".join(lines)


def render_audit_report(package: dict[str, Any]) -> str:
    lines = [
        "# P3 DraftPackage 审计报告",
        "",
        f"- 题目：{package['topic']}",
        f"- 主交付物：`{package['outputs']['docx']}`",
        f"- Markdown 源：`{package['outputs']['markdown']}`",
        f"- 问题清单：`{package['outputs']['issue_list']}`",
        "- 未执行回归：是",
        "- 正式层写回：否",
        "- run id：未创建",
        "",
        "## 边界",
        "",
        "- 未写 `state/product/variable_roles.json`。",
        "- 未写 `state/product/design_spec.json`。",
        "- 未写 `state/product/run_plan.json`。",
        "- 未把半成品论文伪装成完整论文。",
        "",
        "## 证据来源",
        "",
        f"- P2 执行准入：`{package['source_artifacts']['p2_execution_readiness']['path']}`",
        "",
    ]
    return "\n".join(lines)


def write_docx_draft(path: Path, package: dict[str, Any]) -> None:
    document = Document()
    document.add_heading(package["topic"], level=0)
    document.add_heading("摘要", level=1)
    document.add_paragraph(
        "本文拟研究父母受教育水平对子女工资收入的影响。当前版本是半成品论文初稿，能写的研究问题、数据设想和方法路径已经先行整理。"
    )
    red = document.add_paragraph()
    run = red.add_run("【红标】父母教育字段尚未绑定，不能报告回归结果或因果结论。")
    run.font.color.rgb = RGBColor(192, 0, 0)
    document.add_heading("研究问题", level=1)
    document.add_paragraph("本文关注父母受教育水平是否以及如何影响子女工资收入。")
    document.add_heading("数据与变量", level=1)
    document.add_paragraph("当前不能写成正式变量表的部分：父亲教育、母亲教育或父母教育综合指标仍缺真实字段绑定。")
    document.add_heading("识别策略", level=1)
    document.add_paragraph("当前版本只保留方法草案，不执行 IV、DID 或 DML。")
    document.add_heading("问题清单", level=1)
    for issue in package["issues"]:
        document.add_paragraph(f"{issue['id']}：{issue['summary']}；下一步：{issue['next_action']}", style="List Bullet")
    document.save(path)


def load_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def normalize_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item) for item in value if str(item)]
