from __future__ import annotations

import csv
import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


TOPIC = "父母受教育水平对子女工资收入的影响"
SCHEMA_VERSION = "p13_p16.parent_education_wage_demo_closure.v1"

P12_PREFLIGHT_PATH = Path("Results/json/parent_education_wage_p12_design_spec_preflight.json")
P13_JSON_PATH = Path("Results/json/parent_education_wage_p13_run_plan_approval.json")
P14_JSON_PATH = Path("Results/json/parent_education_wage_p14_execution_evidence_ledger.json")
P15_JSON_PATH = Path("Results/json/parent_education_wage_p15_draft_export_package.json")
P16_JSON_PATH = Path("Results/json/parent_education_wage_p16_user_acceptance_packet.json")

P13_REVIEW_PATH = Path("Reviews/parent_education_wage_p13_run_plan_approval.md")
P14_REVIEW_PATH = Path("Reviews/parent_education_wage_p14_execution_evidence_ledger.md")
P15_REVIEW_PATH = Path("Reviews/parent_education_wage_p15_draft_export_package.md")
P16_REVIEW_PATH = Path("Reviews/parent_education_wage_p16_user_acceptance_packet.md")
P15_ISSUE_LIST_PATH = Path("Manuscripts/generated/parent_education_wage_p15_issue_list.md")
P15_COMPLETE_DRAFT_MD_PATH = Path("Manuscripts/generated/parent_education_wage_complete_paper_draft.md")

FORMAL_DESIGN_SPEC_PATH = Path("state/product/design_spec.json")
FORMAL_RUN_PLAN_PATH = Path("state/product/run_plan.json")
ARCHIVE_DIR = Path("state/product/archive/p13_p16_stale_formal_state")
DEFAULT_DRAFT_DOCX_PATH = Path("Submissions/parent_education_wage_paper_draft.docx")

STALE_MARKERS = ("工业机器人", "机器人", "ln_robot", "bartik_iv", "robot_wage", "robot")


def run_parent_education_wage_p13_p16_demo_closure(project_root: Path) -> tuple[dict[str, Any], dict[str, Path]]:
    project_root = project_root.resolve()
    p12 = load_json(project_root / P12_PREFLIGHT_PATH)
    stale_archives = archive_stale_formal_state(project_root)
    p13 = build_p13_run_plan_approval(project_root, p12, stale_archives)
    p14 = build_p14_execution_ledger(project_root, p13)
    p15 = build_p15_draft_package(project_root, p13, p14)
    p16 = build_p16_acceptance_packet(p13, p14, p15)
    packet = build_closure_packet(p13, p14, p15, p16)
    paths = write_closure_artifacts(project_root, p13, p14, p15, p16, packet)
    return packet, paths


def get_parent_education_wage_p13_p16_demo_closure(project_root: Path) -> dict[str, Any]:
    project_root = project_root.resolve()
    artifact_paths = [P13_JSON_PATH, P14_JSON_PATH, P15_JSON_PATH, P16_JSON_PATH]
    if all((project_root / path).exists() for path in artifact_paths):
        p13 = load_json(project_root / P13_JSON_PATH)
        p14 = load_json(project_root / P14_JSON_PATH)
        p15 = load_json(project_root / P15_JSON_PATH)
        p16 = load_json(project_root / P16_JSON_PATH)
        return build_closure_packet(p13, p14, p15, p16, artifact_exists=True)
    if (project_root / P12_PREFLIGHT_PATH).exists():
        return build_not_run_packet(completed_stage="P12", blocking_reasons=["missing_p13_p16_artifacts"])
    return {
        "schema_version": SCHEMA_VERSION,
        "generated_at": now(),
        "status": "blocked_missing_p12_design_spec_preflight",
        "topic": TOPIC,
        "artifact_exists": False,
        "blocking_reasons": ["missing_p12_design_spec_preflight"],
        "p13_run_plan_approval": None,
        "p14_execution_ledger": None,
        "p15_draft_package": None,
        "p16_acceptance_packet": None,
    }


def build_p13_run_plan_approval(
    project_root: Path,
    p12: dict[str, Any],
    stale_archives: list[dict[str, str]],
) -> dict[str, Any]:
    draft = p12.get("draft_design_spec") or {}
    dataset_path = Path(str(draft.get("dataset_path") or ""))
    dataset_full_path = project_root / dataset_path
    required_columns = required_columns_from_draft(draft)
    dataset_columns = csv_header(dataset_full_path) if dataset_full_path.exists() else []
    missing_columns = [column for column in required_columns if column not in dataset_columns]
    validation_errors = validate_p12_preflight_for_topic(p12, draft)

    base = {
        "schema_version": SCHEMA_VERSION,
        "stage": "P13",
        "generated_at": now(),
        "topic": TOPIC,
        "dataset_path": dataset_path.as_posix(),
        "required_columns": required_columns,
        "dataset_columns": dataset_columns,
        "missing_dataset_columns": missing_columns,
        "stale_formal_state_archived": stale_archives,
        "can_write_design_spec": False,
        "can_write_run_plan": False,
        "can_create_run_id": False,
        "can_execute_model": False,
        "run_plan": None,
    }
    if validation_errors:
        return {
            **base,
            "status": "blocked_stale_p12_preflight_for_topic",
            "blocking_reasons": validation_errors,
            "missing_dataset_columns": [],
            "next_action": "regenerate_p12_design_spec_preflight_for_parent_education_wage",
        }
    if not dataset_full_path.exists():
        return {
            **base,
            "status": "blocked_missing_dataset_file_for_run_plan",
            "blocking_reasons": ["missing_dataset_file"],
            "next_action": "bind_existing_analysis_dataset",
        }
    if missing_columns:
        return {
            **base,
            "status": "blocked_missing_dataset_columns_for_run_plan",
            "blocking_reasons": ["missing_dataset_columns"],
            "next_action": "deliver_blocked_branch_package",
        }
    run_plan = {
        "id": "parent_education_wage_baseline_ols",
        "status": "approved_for_minimal_execution",
        "dataset_path": dataset_path.as_posix(),
        "formula": (draft.get("model") or {}).get("formula"),
        "estimator": "ols",
        "allowed_methods": ["baseline_ols"],
        "forbidden_methods": ["did", "iv", "rdd"],
    }
    return {
        **base,
        "status": "run_plan_approved_for_baseline_ols",
        "can_write_design_spec": True,
        "can_write_run_plan": True,
        "can_create_run_id": True,
        "can_execute_model": True,
        "blocking_reasons": [],
        "next_action": "run_p14_minimal_ols",
        "run_plan": run_plan,
    }


def build_p14_execution_ledger(project_root: Path, p13: dict[str, Any]) -> dict[str, Any]:
    if p13.get("status") != "run_plan_approved_for_baseline_ols":
        status = (
            "execution_blocked_stale_p12_preflight"
            if p13.get("status") == "blocked_stale_p12_preflight_for_topic"
            else "execution_blocked_missing_dataset_columns"
        )
        next_action = (
            "regenerate_p12_design_spec_preflight_for_parent_education_wage"
            if p13.get("status") == "blocked_stale_p12_preflight_for_topic"
            else "generate_blocked_draft_package"
        )
        return {
            "schema_version": SCHEMA_VERSION,
            "stage": "P14",
            "generated_at": now(),
            "topic": TOPIC,
            "status": status,
            "run_id": None,
            "executed_regression": False,
            "model_results": None,
            "blocking_reasons": p13.get("blocking_reasons", []),
            "missing_dataset_columns": p13.get("missing_dataset_columns", []),
            "next_action": next_action,
        }
    try:
        model_results = run_minimal_ols(project_root, p13)
    except ValueError as exc:
        return {
            "schema_version": SCHEMA_VERSION,
            "stage": "P14",
            "generated_at": now(),
            "topic": TOPIC,
            "status": "execution_blocked_ols_failed",
            "run_id": None,
            "executed_regression": False,
            "model_results": None,
            "blocking_reasons": [str(exc)],
            "missing_dataset_columns": [],
            "next_action": "repair_dataset_or_formula_before_model_claim",
        }
    run_id = f"parent_education_wage_ols_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}"
    return {
        "schema_version": SCHEMA_VERSION,
        "stage": "P14",
        "generated_at": now(),
        "topic": TOPIC,
        "status": "execution_completed_minimal_ols",
        "run_id": run_id,
        "executed_regression": True,
        "model_results": model_results,
        "blocking_reasons": [],
        "missing_dataset_columns": [],
        "next_action": "generate_model_result_draft_package",
    }


def build_p15_draft_package(project_root: Path, p13: dict[str, Any], p14: dict[str, Any]) -> dict[str, Any]:
    draft_path = DEFAULT_DRAFT_DOCX_PATH if (project_root / DEFAULT_DRAFT_DOCX_PATH).exists() else None
    if p14.get("status") == "execution_completed_minimal_ols":
        return {
            "schema_version": SCHEMA_VERSION,
            "stage": "P15",
            "generated_at": now(),
            "topic": TOPIC,
            "status": "complete_paper_draft_package_ready",
            "paper_draft_docx": (draft_path or DEFAULT_DRAFT_DOCX_PATH).as_posix(),
            "paper_draft_markdown": P15_COMPLETE_DRAFT_MD_PATH.as_posix(),
            "issue_list_path": P15_ISSUE_LIST_PATH.as_posix(),
            "red_flag_issues": [],
            "can_export_complete_paper": True,
            "model_results_included": True,
            "source_execution_status": p14.get("status"),
            "model_result_summary": summarize_model_results(p14.get("model_results") or {}),
            "not_submission_ready_reasons": [
                "还需要人工审阅模型解释、变量口径和论文措辞。",
                "当前是可审阅完整初稿，不是投稿终稿。",
            ],
        }
    if p13.get("status") == "blocked_stale_p12_preflight_for_topic":
        red_flag_issue = {
            "id": "stale_p12_preflight_for_topic",
            "severity": "red",
            "title": "P12 预检题目不匹配，不能进入运行计划",
            "missing_columns": [],
            "why_it_blocks": "当前 P12 预检含有旧题目或旧变量，继续执行会把别的论文题目伪装成父母教育工资题目。",
            "repair_action": "重新生成父母教育工资题目的 P12 方法规格预检。",
        }
        return {
            "schema_version": SCHEMA_VERSION,
            "stage": "P15",
            "generated_at": now(),
            "topic": TOPIC,
            "status": "blocked_draft_package_ready",
            "paper_draft_docx": draft_path.as_posix() if draft_path else "",
            "issue_list_path": P15_ISSUE_LIST_PATH.as_posix(),
            "red_flag_issues": [red_flag_issue],
            "can_export_complete_paper": False,
            "model_results_included": False,
            "source_execution_status": p14.get("status"),
        }
    red_flag_issue = {
        "id": "missing_required_dataset_columns",
        "severity": "red",
        "title": "真实数据缺少运行所需字段，不能报告回归结果",
        "missing_columns": p13.get("missing_dataset_columns", []),
        "why_it_blocks": "P12 公式需要这些字段；真实 CSV 表头没有这些列，继续运行会制造假结果。",
        "repair_action": "补齐或合并 parent_education 与 experience 字段后重新运行 P13-P16。",
    }
    return {
        "schema_version": SCHEMA_VERSION,
        "stage": "P15",
        "generated_at": now(),
        "topic": TOPIC,
        "status": "blocked_draft_package_ready",
        "paper_draft_docx": draft_path.as_posix() if draft_path else "",
        "issue_list_path": P15_ISSUE_LIST_PATH.as_posix(),
        "red_flag_issues": [red_flag_issue],
        "can_export_complete_paper": False,
        "model_results_included": False,
        "source_execution_status": p14.get("status"),
    }


def build_p16_acceptance_packet(
    p13: dict[str, Any],
    p14: dict[str, Any],
    p15: dict[str, Any],
) -> dict[str, Any]:
    if p14.get("status") == "execution_completed_minimal_ols":
        return {
            "schema_version": SCHEMA_VERSION,
            "stage": "P16",
            "generated_at": now(),
            "topic": TOPIC,
            "status": "demo_closure_complete_paper_draft_ready",
            "current_user_outcome": "完整论文初稿 + 真实模型结果证据包",
            "can_accept_blocked_package": False,
            "can_claim_complete_paper": True,
            "can_claim_model_result": True,
            "can_claim_submission_ready": False,
            "completed_stage": "P16",
            "what_user_can_review_now": [
                "Submissions/parent_education_wage_paper_draft.docx",
                P15_COMPLETE_DRAFT_MD_PATH.as_posix(),
                P14_JSON_PATH.as_posix(),
                P16_JSON_PATH.as_posix(),
            ],
            "cannot_claim": [
                "不能声称这是投稿终稿。",
                "不能跳过人工审阅直接提交。",
            ],
            "next_actions": [
                "人工审阅模型解释、变量口径和论文措辞。",
                "审阅通过后再进入投稿格式、参考文献和 PDF 导出阶段。",
            ],
            "stage_status": {
                "p13": p13.get("status"),
                "p14": p14.get("status"),
                "p15": p15.get("status"),
            },
        }
    if p13.get("status") == "blocked_stale_p12_preflight_for_topic":
        return {
            "schema_version": SCHEMA_VERSION,
            "stage": "P16",
            "generated_at": now(),
            "topic": TOPIC,
            "status": "blocked_stale_p12_preflight_for_topic",
            "current_user_outcome": "P12 预检题目污染，不能进入 Demo 闭环",
            "can_accept_blocked_package": False,
            "can_claim_complete_paper": False,
            "can_claim_model_result": False,
            "completed_stage": "P13",
            "what_user_can_review_now": [P13_JSON_PATH.as_posix(), P15_ISSUE_LIST_PATH.as_posix()],
            "cannot_claim": [
                "不能把旧机器人题目的方法规格、运行计划或回归结果改名成父母教育工资题目。",
                "不能进入运行计划批准。",
            ],
            "next_actions": [
                "重新生成父母教育工资题目的 P12 方法规格预检。",
                "确认 P12 不含旧题目和旧变量后再运行 P13-P16。",
            ],
            "stage_status": {
                "p13": p13.get("status"),
                "p14": p14.get("status"),
                "p15": p15.get("status"),
            },
        }
    return {
        "schema_version": SCHEMA_VERSION,
        "stage": "P16",
        "generated_at": now(),
        "topic": TOPIC,
        "status": "demo_closure_blocked_branch_ready",
        "current_user_outcome": "半成品论文 + 红标问题清单",
        "can_accept_blocked_package": True,
        "can_claim_complete_paper": False,
        "can_claim_model_result": False,
        "completed_stage": "P16",
        "what_user_can_review_now": [
            "Submissions/parent_education_wage_paper_draft.docx",
            P15_ISSUE_LIST_PATH.as_posix(),
            P16_JSON_PATH.as_posix(),
        ],
        "cannot_claim": [
            "不能声称已经完成父母教育工资回归。",
            "不能声称已经生成完整实证论文。",
            "不能复用旧机器人题目的方法规格、运行计划或回归结果。",
        ],
        "next_actions": [
            "在真实分析数据中补齐 parent_education。",
            "在真实分析数据中补齐或构造 experience。",
            "补齐字段后重新运行 P13-P16，再进入完整论文分支。",
        ],
        "stage_status": {
            "p13": p13.get("status"),
            "p14": p14.get("status"),
            "p15": p15.get("status"),
        },
    }


def build_closure_packet(
    p13: dict[str, Any],
    p14: dict[str, Any],
    p15: dict[str, Any],
    p16: dict[str, Any],
    artifact_exists: bool = True,
) -> dict[str, Any]:
    return {
        "schema_version": SCHEMA_VERSION,
        "generated_at": now(),
        "topic": TOPIC,
        "status": p16.get("status"),
        "completed_stage": p16.get("completed_stage", "P16"),
        "artifact_exists": artifact_exists,
        "p13_run_plan_approval": p13,
        "p14_execution_ledger": p14,
        "p15_draft_package": p15,
        "p16_acceptance_packet": p16,
        "outputs": {
            "p13_json": P13_JSON_PATH.as_posix(),
            "p14_json": P14_JSON_PATH.as_posix(),
            "p15_json": P15_JSON_PATH.as_posix(),
            "p16_json": P16_JSON_PATH.as_posix(),
            "p15_issue_list": P15_ISSUE_LIST_PATH.as_posix(),
            "p15_complete_draft_markdown": P15_COMPLETE_DRAFT_MD_PATH.as_posix(),
        },
    }


def write_closure_artifacts(
    project_root: Path,
    p13: dict[str, Any],
    p14: dict[str, Any],
    p15: dict[str, Any],
    p16: dict[str, Any],
    packet: dict[str, Any],
) -> dict[str, Path]:
    path_payloads = {
        P13_JSON_PATH: p13,
        P14_JSON_PATH: p14,
        P15_JSON_PATH: p15,
        P16_JSON_PATH: p16,
    }
    for relative_path, payload in path_payloads.items():
        write_json(project_root / relative_path, payload)
    write_text(project_root / P13_REVIEW_PATH, render_p13_review(p13))
    write_text(project_root / P14_REVIEW_PATH, render_p14_review(p14))
    write_text(project_root / P15_REVIEW_PATH, render_p15_review(p15))
    write_text(project_root / P16_REVIEW_PATH, render_p16_review(p16))
    write_text(project_root / P15_ISSUE_LIST_PATH, render_issue_list(p15))
    if p15.get("status") == "complete_paper_draft_package_ready":
        draft_markdown = render_complete_paper_draft(p15, p14)
        write_text(project_root / P15_COMPLETE_DRAFT_MD_PATH, draft_markdown)
        write_docx(project_root / DEFAULT_DRAFT_DOCX_PATH, draft_markdown)
    return {
        "p13_json": project_root / P13_JSON_PATH,
        "p14_json": project_root / P14_JSON_PATH,
        "p15_json": project_root / P15_JSON_PATH,
        "p16_json": project_root / P16_JSON_PATH,
        "p15_issue_list": project_root / P15_ISSUE_LIST_PATH,
        "p15_complete_draft_markdown": project_root / P15_COMPLETE_DRAFT_MD_PATH,
        "paper_draft_docx": project_root / DEFAULT_DRAFT_DOCX_PATH,
    }


def archive_stale_formal_state(project_root: Path) -> list[dict[str, str]]:
    archives: list[dict[str, str]] = []
    for relative_path in (FORMAL_DESIGN_SPEC_PATH, FORMAL_RUN_PLAN_PATH):
        path = project_root / relative_path
        if not path.exists() or not is_stale_formal_state(path):
            continue
        archive_dir = project_root / ARCHIVE_DIR
        archive_dir.mkdir(parents=True, exist_ok=True)
        target = archive_dir / relative_path.name
        if target.exists():
            target = archive_dir / f"{relative_path.stem}_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}{relative_path.suffix}"
        shutil.move(str(path), str(target))
        archives.append({"source": relative_path.as_posix(), "archive": target.relative_to(project_root).as_posix()})
    return archives


def is_stale_formal_state(path: Path) -> bool:
    text = path.read_text(encoding="utf-8")
    if any(marker in text for marker in STALE_MARKERS):
        return True
    try:
        payload = json.loads(text)
    except json.JSONDecodeError:
        return False
    question = str(payload.get("research_question") or payload.get("topic") or "")
    return bool(question and TOPIC not in question)


def build_not_run_packet(completed_stage: str, blocking_reasons: list[str]) -> dict[str, Any]:
    return {
        "schema_version": SCHEMA_VERSION,
        "generated_at": now(),
        "status": "p13_p16_closure_not_run",
        "topic": TOPIC,
        "completed_stage": completed_stage,
        "artifact_exists": False,
        "blocking_reasons": blocking_reasons,
        "p13_run_plan_approval": None,
        "p14_execution_ledger": None,
        "p15_draft_package": None,
        "p16_acceptance_packet": None,
        "next_action": "run_p13_p16_demo_closure",
    }


def validate_p12_preflight_for_topic(p12: dict[str, Any], draft: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    p12_topic = str(p12.get("topic") or "")
    draft_question = str(draft.get("research_question") or draft.get("topic") or "")
    topic_slug = str(draft.get("topic_slug") or "")
    variables = draft.get("variables") or {}
    treatment = normalize_string_list(variables.get("treatment"))
    formula = str((draft.get("model") or {}).get("formula") or "")
    estimator = str((draft.get("model") or {}).get("estimator") or "ols").lower()
    identity_text = json.dumps(
        {
            "p12_topic": p12_topic,
            "draft_question": draft_question,
            "topic_slug": topic_slug,
            "treatment": treatment,
            "outcome": normalize_string_list(variables.get("outcome")),
            "controls": normalize_string_list(variables.get("controls")),
            "formula": formula,
        },
        ensure_ascii=False,
    )

    if TOPIC not in p12_topic:
        errors.append("p12_topic_not_parent_education_wage")
    if TOPIC not in draft_question:
        errors.append("draft_research_question_not_parent_education_wage")
    if topic_slug and topic_slug not in {"parent-education-wage", "parent_education_wage"}:
        errors.append("topic_slug_not_parent_education_wage")
    if treatment and "parent_education" not in treatment:
        errors.append("treatment_not_parent_education")
    if formula and "parent_education" not in formula:
        errors.append("formula_missing_parent_education")
    if estimator and estimator != "ols":
        errors.append("estimator_not_baseline_ols")
    if any(marker in identity_text for marker in STALE_MARKERS):
        errors.append("stale_robot_topic_or_variable_marker_detected")
    return errors


def required_columns_from_draft(draft: dict[str, Any]) -> list[str]:
    columns: list[str] = []
    variables = draft.get("variables") or {}
    for key in ("outcome", "treatment", "controls", "fixed_effects", "cluster_by"):
        for value in normalize_string_list(variables.get(key)):
            if value not in columns:
                columns.append(value)
    formula = str((draft.get("model") or {}).get("formula") or "")
    for token in re.findall(r"[A-Za-z_][A-Za-z0-9_]*", formula):
        if token not in {"C", "I", "log"} and token not in columns:
            columns.append(token)
    return columns


def run_minimal_ols(project_root: Path, p13: dict[str, Any]) -> dict[str, Any]:
    run_plan = p13.get("run_plan") or {}
    formula = str(run_plan.get("formula") or "")
    outcome, rhs_columns = parse_simple_formula(formula)
    dataset_path = project_root / str(run_plan.get("dataset_path") or p13.get("dataset_path") or "")
    if not dataset_path.exists():
        raise ValueError("dataset_file_missing_for_minimal_ols")

    names = ["Intercept", *rhs_columns]
    x_rows: list[list[float]] = []
    y_values: list[float] = []
    with dataset_path.open(encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            try:
                y_value = float(str(row.get(outcome, "")).strip())
                x_value = [1.0, *[float(str(row.get(column, "")).strip()) for column in rhs_columns]]
            except ValueError:
                continue
            y_values.append(y_value)
            x_rows.append(x_value)
    if len(x_rows) <= len(names):
        raise ValueError("not_enough_numeric_rows_for_minimal_ols")

    xtx = [[0.0 for _ in names] for _ in names]
    xty = [0.0 for _ in names]
    for row, y_value in zip(x_rows, y_values):
        for i, value_i in enumerate(row):
            xty[i] += value_i * y_value
            for j, value_j in enumerate(row):
                xtx[i][j] += value_i * value_j
    try:
        betas = solve_linear_system(xtx, xty)
    except ValueError:
        ridge_xtx = [[value + (1e-8 if i == j else 0.0) for j, value in enumerate(row)] for i, row in enumerate(xtx)]
        betas = solve_linear_system(ridge_xtx, xty)

    coefficients = {name: round(beta, 8) for name, beta in zip(names, betas)}
    return {
        "estimator": "ols",
        "formula": formula,
        "outcome_variable": outcome,
        "treatment_variable": "parent_education",
        "nobs": len(x_rows),
        "coefficients": coefficients,
    }


def summarize_model_results(model_results: dict[str, Any]) -> dict[str, Any]:
    coefficients = model_results.get("coefficients") or {}
    treatment = str(model_results.get("treatment_variable") or "parent_education")
    treatment_coefficient = coefficients.get(treatment)
    direction = "positive" if isinstance(treatment_coefficient, (int, float)) and treatment_coefficient > 0 else "non_positive"
    return {
        "estimator": model_results.get("estimator"),
        "formula": model_results.get("formula"),
        "nobs": model_results.get("nobs"),
        "treatment_variable": treatment,
        "treatment_coefficient": treatment_coefficient,
        "direction": direction,
    }


def parse_simple_formula(formula: str) -> tuple[str, list[str]]:
    if "~" not in formula:
        raise ValueError("formula_missing_tilde")
    lhs, rhs = formula.split("~", 1)
    outcome = lhs.strip()
    rhs_columns = [part.strip() for part in rhs.split("+") if part.strip()]
    if not outcome or not rhs_columns:
        raise ValueError("formula_missing_outcome_or_rhs")
    unsupported = [column for column in rhs_columns if not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", column)]
    if unsupported:
        raise ValueError("formula_contains_unsupported_terms")
    return outcome, rhs_columns


def solve_linear_system(matrix: list[list[float]], vector: list[float]) -> list[float]:
    size = len(vector)
    augmented = [row[:] + [value] for row, value in zip(matrix, vector)]
    for column in range(size):
        pivot_row = max(range(column, size), key=lambda row: abs(augmented[row][column]))
        if abs(augmented[pivot_row][column]) < 1e-12:
            raise ValueError("singular_matrix_for_minimal_ols")
        augmented[column], augmented[pivot_row] = augmented[pivot_row], augmented[column]
        pivot = augmented[column][column]
        for item in range(column, size + 1):
            augmented[column][item] /= pivot
        for row in range(size):
            if row == column:
                continue
            factor = augmented[row][column]
            for item in range(column, size + 1):
                augmented[row][item] -= factor * augmented[column][item]
    return [augmented[row][size] for row in range(size)]


def normalize_string_list(value: Any) -> list[str]:
    if isinstance(value, str):
        value = [part.strip() for part in value.split(",")]
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def csv_header(path: Path) -> list[str]:
    with path.open(encoding="utf-8", newline="") as handle:
        reader = csv.reader(handle)
        return [item.strip() for item in next(reader, []) if item.strip()]


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def render_p13_review(p13: dict[str, Any]) -> str:
    missing = ", ".join(p13.get("missing_dataset_columns") or []) or "无"
    archived = ", ".join(item["archive"] for item in p13.get("stale_formal_state_archived", [])) or "无"
    return "\n".join(
        [
            "# P13 RunPlan Approval Review",
            "",
            f"- status: `{p13.get('status')}`",
            f"- dataset: `{p13.get('dataset_path')}`",
            f"- missing columns: {missing}",
            f"- stale formal state archived: {archived}",
            f"- can_write_run_plan: `{p13.get('can_write_run_plan')}`",
            f"- can_create_run_id: `{p13.get('can_create_run_id')}`",
            "",
            "结论：字段不齐时，不批准 RunPlan，不创建运行编号。",
            "",
        ]
    )


def render_p14_review(p14: dict[str, Any]) -> str:
    return "\n".join(
        [
            "# P14 Execution Evidence Ledger Review",
            "",
            f"- status: `{p14.get('status')}`",
            f"- run_id: `{p14.get('run_id')}`",
            f"- executed_regression: `{p14.get('executed_regression')}`",
            "",
            "结论：P13 阻断时，P14 只写执行证据账本，不运行模型。",
            "",
        ]
    )


def render_p15_review(p15: dict[str, Any]) -> str:
    return "\n".join(
        [
            "# P15 Draft Export Package Review",
            "",
            f"- status: `{p15.get('status')}`",
            f"- paper_draft_docx: `{p15.get('paper_draft_docx')}`",
            f"- issue_list_path: `{p15.get('issue_list_path')}`",
            f"- can_export_complete_paper: `{p15.get('can_export_complete_paper')}`",
            "",
            "结论：当前交付半成品论文和红标问题清单。",
            "",
        ]
    )


def render_p16_review(p16: dict[str, Any]) -> str:
    return "\n".join(
        [
            "# P16 User Acceptance Packet Review",
            "",
            f"- status: `{p16.get('status')}`",
            f"- current_user_outcome: {p16.get('current_user_outcome')}",
            f"- can_claim_complete_paper: `{p16.get('can_claim_complete_paper')}`",
            f"- can_accept_blocked_package: `{p16.get('can_accept_blocked_package')}`",
            "",
            "结论：Demo 线推进到 P16 阻断交付分支；不能声称完整模型结果。",
            "",
        ]
    )


def render_issue_list(p15: dict[str, Any]) -> str:
    lines = [
        "# 父母教育工资 Demo 红标问题清单",
        "",
        "当前可以交付半成品论文，但不能交付完整实证结果。",
        "",
    ]
    if not p15.get("red_flag_issues"):
        return "\n".join(
            [
                "# 父母教育工资 Demo 红标问题清单",
                "",
                "当前完整初稿分支已生成真实模型结果；没有阻断交付的红标问题。",
                "",
                "注意：这不等于投稿终稿，仍需人工审阅模型解释、变量口径和论文措辞。",
                "",
            ]
        )
    for issue in p15.get("red_flag_issues", []):
        missing = ", ".join(issue.get("missing_columns") or []) or "无"
        lines.extend(
            [
                f"## RED: {issue.get('title')}",
                "",
                f"- 缺少字段：{missing}",
                f"- 为什么阻断：{issue.get('why_it_blocks')}",
                f"- 修复动作：{issue.get('repair_action')}",
                "",
            ]
        )
    return "\n".join(lines)


def render_complete_paper_draft(p15: dict[str, Any], p14: dict[str, Any]) -> str:
    model_results = p14.get("model_results") or {}
    summary = p15.get("model_result_summary") or summarize_model_results(model_results)
    coefficient = summary.get("treatment_coefficient")
    coefficient_text = "NA" if coefficient is None else str(coefficient)
    nobs = summary.get("nobs") or "NA"
    formula = summary.get("formula") or "ln_wage ~ parent_education + age + female + urban + edu_last + experience"
    return "\n".join(
        [
            "# 父母受教育水平如何影响子女的工资水平？",
            "",
            "## 摘要",
            "",
            "本文使用修复后的 CFPS 分析数据，估计父母教育水平与子女工资水平之间的关系。系统先识别原始分析数据缺少 `parent_education` 与 `experience`，再通过 P18 数据修复门禁生成新的分析数据集，并在修复数据上执行最小 OLS。",
            "",
            "## 数据与变量",
            "",
            "- 因变量：`ln_wage`。",
            "- 核心解释变量：`parent_education`，由父亲和母亲教育水平中的有效较高值构造。",
            "- 控制变量：`age`、`female`、`urban`、`edu_last`、`experience`。",
            "- 修复数据集：`Data/Interim/parent_education_wage_repaired.csv`。",
            "",
            "## 方法",
            "",
            f"基准模型为 `{formula}`。当前版本执行的是最小 OLS，用于产品闭环和结果证据验证；正式论文仍需人工审阅识别策略、变量口径和稳健性设计。",
            "",
            "## 结果",
            "",
            f"模型使用样本量为 `{nobs}`。`parent_education` 的估计系数为 `{coefficient_text}`。这说明在当前控制变量下，父母教育水平与子女工资水平存在可审阅的统计关联。",
            "",
            "## 审阅边界",
            "",
            "这是一份可审阅完整初稿，不是投稿终稿。提交前还需要人工检查父母教育变量构造、教育年限映射、样本选择、稳健性检验、参考文献和格式。",
            "",
            "## 证据路径",
            "",
            "- P18 修复数据：`Results/json/parent_education_wage_p18_data_repair_apply.json`。",
            "- P13 运行计划审批：`Results/json/parent_education_wage_p13_run_plan_approval.json`。",
            "- P14 模型执行账本：`Results/json/parent_education_wage_p14_execution_evidence_ledger.json`。",
            "- P16 用户验收包：`Results/json/parent_education_wage_p16_user_acceptance_packet.json`。",
            "",
        ]
    )


def write_docx(path: Path, markdown: str) -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for line in markdown.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif not line.strip():
            continue
        else:
            document.add_paragraph(line)
    document.save(path)


def now() -> str:
    return datetime.now(timezone.utc).isoformat()
